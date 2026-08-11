"""
Knowledge Ingestion for SpendWise-AI — Improved Pipeline
=========================================================
Fixes vs original:
  1. Section-aware splitting — splits on ===...=== headers BEFORE size-based chunking,
     so chunks never bleed across semantic section boundaries.
  2. Rich metadata — section_title, category, persona_relevance, source_file, chunk_hash.
  3. Content-hash IDs — MD5(text[:64]) → idempotent re-ingestion without duplicates.
  4. upsert() instead of add() — safe to re-run.
  5. embed_documents() instead of embed_query() — correct asymmetric embedding side.
  6. Smaller chunk_size=600 within sections — tighter, more focused embedding signal.
  7. Proper logging — no silent except:pass.
  8. Docx structure-aware loading via python-docx (headings preserved).
"""

import os
import re
import hashlib
import logging
from pathlib import Path

import chromadb
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_huggingface import HuggingFaceEmbeddings


load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
logger = logging.getLogger(__name__)

# =====================================================
# CONFIG
# =====================================================
KNOWLEDGE_DIR = "finance_knowledge_base"
CHROMA_PATH   = "chroma_db"
CHUNK_SIZE    = 600    # smaller than original 1000 — keeps chunks focused
CHUNK_OVERLAP = 80     # enough for sentence continuity

# ── Category & persona inference from filename stems ──────────────────────
_FILENAME_CATEGORY = {
    "tax":        "tax",
    "investment": "investment",
    "debt":       "debt",
    "emergency":  "insurance",
    "insurance":  "insurance",
    "expense":    "savings",
    "income":     "income",
}

_FILENAME_PERSONA = {
    "tax":        ["salaried", "retiree"],
    "investment": ["all"],
    "debt":       ["salaried", "student"],
    "emergency":  ["all"],
    "insurance":  ["all"],
    "expense":    ["all"],
    "income":     ["salaried", "student"],
}

def _infer_metadata_from_filename(filepath: str):
    stem = Path(filepath).stem.lower()
    category = "general"
    personas = ["all"]
    for kw, cat in _FILENAME_CATEGORY.items():
        if kw in stem:
            category = cat
            personas = _FILENAME_PERSONA.get(kw, ["all"])
            break
    return category, ",".join(personas)



# =====================================================
# EMBEDDING MODEL
# =====================================================

embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)


# =====================================================
# CHROMA CLIENT
# =====================================================
client = chromadb.PersistentClient(path=CHROMA_PATH)
knowledge_collection = client.get_or_create_collection(name="financial_knowledge")
user_collection       = client.get_or_create_collection(name="user_profiles")


# =====================================================
# SECTION-AWARE SPLITTING
# =====================================================
# Pattern matches ===...=== header blocks (used in all .txt files)
_SECTION_HEADER_RE = re.compile(
    r"={3,}[\r\n]+(.*?)[\r\n]+=={3,}",
    re.MULTILINE
)

def _split_into_sections(text: str) -> list[tuple[str, str]]:
    """
    Splits document text on ===...=== header blocks.
    Returns list of (section_title, section_body) tuples.
    Falls back to a single section if no headers found.
    """
    parts = _SECTION_HEADER_RE.split(text)
    # parts alternates: [pre-text, title1, body1, title2, body2, ...]
    if len(parts) <= 1:
        return [("General", text.strip())]

    sections = []
    # First part before any header
    if parts[0].strip():
        sections.append(("Introduction", parts[0].strip()))
    # Iterate title/body pairs
    for i in range(1, len(parts) - 1, 2):
        title = parts[i].strip()
        body  = parts[i + 1].strip() if i + 1 < len(parts) else ""
        if body:
            sections.append((title, body))
    return sections


def _load_txt(filepath: str) -> list[tuple[str, str]]:
    """Load .txt and split into (section_title, body) tuples."""
    with open(filepath, encoding="utf-8", errors="replace") as f:
        text = f.read()
    return _split_into_sections(text)


def _load_docx(filepath: str) -> list[tuple[str, str]]:
    """
    Load .docx using python-docx to preserve heading → paragraph structure.
    Falls back to plain-text extraction if python-docx unavailable.
    """
    try:
        import docx  # python-docx
        doc = docx.Document(filepath)
        sections, current_title, current_body = [], "Introduction", []
        for para in doc.paragraphs:
            style = para.style.name.lower() if para.style else ""
            text  = para.text.strip()
            if not text:
                continue
            if "heading" in style:
                if current_body:
                    sections.append((current_title, " ".join(current_body)))
                current_title = text
                current_body  = []
            else:
                current_body.append(text)
        if current_body:
            sections.append((current_title, " ".join(current_body)))
        return sections if sections else [("General", " ".join(p.text for p in doc.paragraphs))]
    except ImportError:
        #https://featurepilot-openai.openai.azure.com/openai/deployments/text-embedding-3-small/embeddings?api-version=2023-05-15
        logger.warning("python-docx not installed; falling back to plain-text for %s", filepath)
        try:
            import docx2txt
            text = docx2txt.process(filepath)
            return _split_into_sections(text)
        except Exception as e:
            logger.error("Failed to load %s: %s", filepath, e)
            return []


def _load_pdf(filepath: str) -> list[tuple[str, str]]:
    """Load PDF pages; treat each page as a section."""
    try:
        from langchain_community.document_loaders import PyPDFLoader
        docs = PyPDFLoader(filepath).load()
        return [(f"Page {d.metadata.get('page', i+1)}", d.page_content.strip())
                for i, d in enumerate(docs) if d.page_content.strip()]
    except Exception as e:
        logger.error("Failed to load PDF %s: %s", filepath, e)
        return []


def load_all_sections() -> list[dict]:
    """
    Load all documents and return a flat list of section dicts:
      {text, section_title, source_file, category, persona_relevance}
    """
    base = Path(KNOWLEDGE_DIR)
    all_sections = []
    found = list(base.glob("**/*.txt")) + list(base.glob("**/*.docx")) + list(base.glob("**/*.pdf"))

    if not found:
        logger.warning("No documents found in %s", KNOWLEDGE_DIR)
        return []

    for fp in found:
        category, personas = _infer_metadata_from_filename(str(fp))
        ext = fp.suffix.lower()
        try:
            if ext == ".txt":
                raw_sections = _load_txt(str(fp))
            elif ext == ".docx":
                raw_sections = _load_docx(str(fp))
            elif ext == ".pdf":
                raw_sections = _load_pdf(str(fp))
            else:
                continue
        except Exception as e:
            logger.error("Loader failed for %s: %s", fp, e)
            continue

        for title, body in raw_sections:
            if body.strip():
                all_sections.append({
                    "text":              body.strip(),
                    "section_title":     title,
                    "source_file":       fp.name,
                    "category":          category,
                    "persona_relevance": personas,
                })

    logger.info("Loaded %d sections from %d files", len(all_sections), len(found))
    return all_sections


# =====================================================
# SUB-CHUNKING WITHIN SECTIONS
# =====================================================
_sub_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
)

def chunk_sections(sections: list[dict]) -> list[dict]:
    """
    For sections longer than CHUNK_SIZE, apply RecursiveCharacterTextSplitter.
    Each sub-chunk inherits the section's metadata plus a sub-index.
    Returns flat list of chunk dicts ready for ingestion.
    """
    chunks = []
    for sec in sections:
        text = sec["text"]
        if len(text) <= CHUNK_SIZE:
            chunks.append({**sec, "sub_index": 0})
        else:
            sub_texts = _sub_splitter.split_text(text)
            for i, sub in enumerate(sub_texts):
                if sub.strip():
                    chunks.append({**sec, "text": sub.strip(), "sub_index": i})
    logger.info("Created %d chunks from %d sections", len(chunks), len(sections))
    return chunks


# =====================================================
# CONTENT-HASH ID (idempotent)
# =====================================================
def _chunk_id(chunk: dict) -> str:
    """Stable ID = MD5 of (source_file + section_title + first 80 chars of text)."""
    key = f"{chunk['source_file']}|{chunk['section_title']}|{chunk['text'][:80]}"
    return hashlib.md5(key.encode()).hexdigest()


# =====================================================
# INGEST KNOWLEDGE BASE
# =====================================================
def ingest_knowledge():
    sections = load_all_sections()
    if not sections:
        logger.error("No documents found. Aborting ingestion.")
        return

    chunks = chunk_sections(sections)

    ids, texts, metadatas = [], [], []
    for chunk in chunks:
        ids.append(_chunk_id(chunk))
        texts.append(chunk["text"])
        metadatas.append({
            "section_title":     chunk["section_title"],
            "source_file":       chunk["source_file"],
            "category":          chunk["category"],
            "persona_relevance": chunk["persona_relevance"],
            "sub_index":         str(chunk["sub_index"]),
        })

    # Batch embed using embed_documents (correct side for passage encoding)
    logger.info("Embedding %d chunks...", len(texts))
    embeddings = embedding_model.embed_documents(texts)

    # Upsert in batches of 100 (ChromaDB recommended)
    BATCH = 100
    total_upserted = 0
    for start in range(0, len(ids), BATCH):
        end = start + BATCH
        knowledge_collection.upsert(
            ids=ids[start:end],
            documents=texts[start:end],
            metadatas=metadatas[start:end],
            embeddings=embeddings[start:end],
        )
        total_upserted += len(ids[start:end])
        logger.info("Upserted batch %d-%d", start, end)

    logger.info("✅ Ingestion complete: %d chunks upserted to ChromaDB", total_upserted)

    # Print category summary
    from collections import Counter
    cats = Counter(m["category"] for m in metadatas)
    for cat, count in cats.most_common():
        logger.info("  Category %-12s → %d chunks", cat, count)


if __name__ == "__main__":
    ingest_knowledge()





